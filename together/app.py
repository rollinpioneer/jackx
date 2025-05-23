from flask import Flask, render_template, request, jsonify, url_for, send_file
import subprocess
import os
import random
import re
import json
import sys
from evaluate_summary import evaluate_from_file, evaluate_summary
import docx2txt  # 用于解析Word文档
from openai import OpenAI  # 导入OpenAI库用于调用DeepSeek API
import tempfile
from docx import Document  # 用于创建Word文档
import uuid
import time
import threading
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import torch

app = Flask(__name__)
app.static_folder = 'static'

# 文章数据的全局变量
current_data = {
    "news": {"article": "", "reference": "", "index": -1},
    "media": {"article": "", "reference": "", "index": -1},
    "encyclopedia": {"article": "", "reference": "", "index": -1},
    "custom": {"article": "", "index": -1}
}

# 允许上传的文件类型
ALLOWED_EXTENSIONS = {'docx', 'doc'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# 处理文本中空格的函数
def clean_sentence_spaces(text):
    if not text:
        return ""
    sentences = re.split('([。！？!?])', text)
    result = ""
    for i in range(0, len(sentences), 2):
        if i < len(sentences):
            current = sentences[i].strip()
            if current:
                result += current
                if i+1 < len(sentences):
                    result += sentences[i+1]
    return result

# 处理空白的函数 - 供批量摘要使用
WHITESPACE_HANDLER = lambda k: re.sub('\s+', ' ', re.sub('\n+', ' ', k.strip()))

# 加载Transformer模型供批量摘要使用 - 延迟加载模式
transformer_model = None
transformer_tokenizer = None

def load_mt5_model():
    global transformer_model, transformer_tokenizer
    if transformer_model is None or transformer_tokenizer is None:
        print("首次加载摘要模型...")
        model_path = r"C:\Users\jackx\.cache\huggingface\hub\models--csebuetnlp--mT5_multilingual_XLSum\snapshots\2437a524effdbadc327ced84595508f1e32025b3"
        transformer_tokenizer = AutoTokenizer.from_pretrained(model_path)
        transformer_model = AutoModelForSeq2SeqLM.from_pretrained(model_path)
        
        # 检查GPU可用性并将模型加载到GPU
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        transformer_model = transformer_model.to(device)
        print(f"摘要模型已加载到 {device}")
    return transformer_model, transformer_tokenizer

# 生成单个文本的摘要
def generate_summary_for_text(text):
    # 确保模型已加载
    model, tokenizer = load_mt5_model()
    device = next(model.parameters()).device
    
    # 清理文本
    text = clean_sentence_spaces(text)
    
    # 生成摘要
    input_ids = tokenizer(
        [WHITESPACE_HANDLER(text)],
        return_tensors="pt",
        padding="max_length",
        truncation=True,
        max_length=512
    )["input_ids"].to(device)

    output_ids = model.generate(
        input_ids=input_ids,
        max_length=84,
        no_repeat_ngram_size=2,
        num_beams=4
    )[0]

    summary = tokenizer.decode(
        output_ids,
        skip_special_tokens=True,
        clean_up_tokenization_spaces=False
    )
    
    # 清理生成摘要中每个句子前的空格
    return clean_sentence_spaces(summary)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/news')
def news():
    return render_template('news.html')

@app.route('/media')
def media():
    return render_template('media.html')

@app.route('/encyclopedia')
def encyclopedia():
    return render_template('encyclopedia.html')

@app.route('/custom')
def custom():
    return render_template('custom.html')

# 批量文档摘要生成
@app.route('/batch_summarize', methods=['POST'])
def batch_summarize():
    try:
        # 检查是否有文件被上传
        if 'files[]' not in request.files:
            return jsonify({"status": "error", "message": "没有选择文件"}), 400
        
        files = request.files.getlist('files[]')
        if not files or len(files) == 0 or files[0].filename == '':
            return jsonify({"status": "error", "message": "没有选择文件"}), 400
        
        # 验证所有文件类型
        for file in files:
            if not allowed_file(file.filename):
                return jsonify({
                    "status": "error", 
                    "message": f"文件 {file.filename} 不是有效的Word文档 (.doc, .docx)"
                }), 400
        
        # 预先加载模型
        load_mt5_model()
        
        # 创建临时目录存储上传的文件
        temp_dir = tempfile.mkdtemp()
        
        # 创建结果文档
        result_doc = Document()
        result_doc.add_heading('批量文章摘要生成结果', 0)
        
        # 处理每个文件
        for file in files:
            # 保存上传的文件到临时路径
            temp_file_path = os.path.join(temp_dir, file.filename)
            file.save(temp_file_path)
            
            try:
                # 提取Word文档内容
                text_content = docx2txt.process(temp_file_path)
                text_content = clean_sentence_spaces(text_content)
                
                # 生成摘要
                summary = generate_summary_for_text(text_content)
                
                # 添加到结果文档
                result_doc.add_heading(f'文件: {file.filename}', level=1)
                
                # 添加原文和摘要
                result_doc.add_heading('原文:', level=2)
                # 截取前500个字符作为预览，避免文档过大
                preview_text = text_content[:500] + ('...' if len(text_content) > 500 else '')
                result_doc.add_paragraph(preview_text)
                
                result_doc.add_heading('摘要:', level=2)
                result_doc.add_paragraph(summary)
                
                # 添加分隔线
                result_doc.add_paragraph('─' * 50)
            
            except Exception as e:
                # 如果处理某个文件出错，添加错误信息到结果文档
                result_doc.add_heading(f'文件: {file.filename}', level=1)
                result_doc.add_paragraph(f'处理出错: {str(e)}', style='Intense Quote')
                result_doc.add_paragraph('─' * 50)
            
            finally:
                # 删除临时文件
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
        
        # 保存结果文档到临时文件
        result_file_path = os.path.join(temp_dir, f'摘要汇总_{uuid.uuid4().hex}.docx')
        result_doc.save(result_file_path)
        
        # 返回文档文件
        return send_file(
            result_file_path,
            as_attachment=True,
            download_name='摘要汇总.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        return jsonify({"status": "error", "message": f"批量处理出错: {str(e)}\n{error_details}"}), 500

# 新增：处理自定义文本处理请求
@app.route('/process_custom_text', methods=['POST'])
def process_custom_text():
    try:
        # 获取请求数据
        data = request.json
        text = data.get('text', '')
        mode = data.get('mode', 'summary')
        prompt = data.get('prompt', '')
        
        if not text:
            return jsonify({"status": "error", "message": "未提供文本内容"})
        
        # 根据不同模式设置系统提示和用户提示
        system_prompts = {
            'summary': "你是一个专业的中文文本摘要助手，只输出简明摘要。",
            'qa': "你是一个专业的文本问答助手，请基于提供的文本回答问题，不要编造信息。",
            'analysis': "你是一个专业的文本分析师，请对提供的文本进行全面分析，包括主题、结构、观点、语言特点等。",
            'key_points': "你是一个专业的文本提炼专家，请从提供的文本中提取5-10个关键点，用简洁的要点形式列出。"
        }
        
        user_prompts = {
            'summary': f"请帮我对下面这段文本生成简明摘要：\n{text}",
            'qa': f"文本内容：\n{text}\n\n问题：{prompt}",
            'analysis': f"请分析以下文本的内容特点、主要观点和写作手法：\n{text}",
            'key_points': f"请从以下文本中提取核心关键点，以要点列表形式呈现：\n{text}"
        }
        
        # 选择对应的提示
        system_prompt = system_prompts.get(mode, system_prompts['summary'])
        user_prompt = user_prompts.get(mode, user_prompts['summary'])
        
        # 调用DeepSeek API
        result = generate_with_deepseek(system_prompt, user_prompt)
        
        return jsonify({
            "status": "success",
            "result": result
        })
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        return jsonify({"status": "error", "message": f"处理文本时出错: {str(e)}\n{error_details}"})

# 使用DeepSeek API生成内容
def generate_with_deepseek(system_prompt, user_prompt):
    try:
        client = OpenAI(api_key="sk-a17fda1fa3bf42e186d7e3868c88f3a8", base_url="https://api.deepseek.com")
        
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            stream=False
        )
        
        # 获取生成的文本
        generated_text = response.choices[0].message.content.strip()
        return generated_text
    except Exception as e:
        print(f"调用DeepSeek API时出错: {e}")
        raise e

# 新增：处理Word文件上传
@app.route('/upload_word', methods=['POST'])
def upload_word():
    try:
        # 检查是否有文件被上传
        if 'wordFile' not in request.files:
            return jsonify({"status": "error", "message": "没有选择文件"})
        
        file = request.files['wordFile']
        
        # 如果用户没有选择文件，浏览器也会提交一个没有文件名的空文件部分
        if file.filename == '':
            return jsonify({"status": "error", "message": "没有选择文件"})
        
        # 检查文件类型
        if file and allowed_file(file.filename):
            # 保存上传的文件到临时路径
            temp_path = os.path.join(os.path.dirname(__file__), "temp_word_file.docx")
            file.save(temp_path)
            
            # 提取Word文档内容
            try:
                text_content = docx2txt.process(temp_path)
                text_content = clean_sentence_spaces(text_content)
                
                # 提取文章类型
                article_type = request.form.get('articleType', 'news')
                
                # 根据文章类型保存到相应的当前文件
                if article_type == 'news':
                    current_file = os.path.join(os.path.dirname(__file__), "current_news.txt")
                    current_data["news"]["article"] = text_content
                    current_data["news"]["reference"] = ""
                    current_data["news"]["index"] = -1
                elif article_type == 'media':
                    current_file = os.path.join(os.path.dirname(__file__), "current_media.txt")
                    current_data["media"]["article"] = text_content
                    current_data["media"]["reference"] = ""
                    current_data["media"]["index"] = -1
                elif article_type == 'encyclopedia':
                    current_file = os.path.join(os.path.dirname(__file__), "current_encyclopedia.txt")
                    current_data["encyclopedia"]["article"] = text_content
                    current_data["encyclopedia"]["reference"] = ""
                    current_data["encyclopedia"]["index"] = -1
                elif article_type == 'custom':
                    current_file = os.path.join(os.path.dirname(__file__), "current_custom.txt")
                    current_data["custom"]["article"] = text_content
                    current_data["custom"]["index"] = -1
                
                # 将提取的文本保存到当前文件
                with open(current_file, "w", encoding="utf-8") as f:
                    f.write(f"ARTICLE:{text_content}\n")
                    if article_type != 'custom' and article_type != 'media':
                        f.write(f"REFERENCE:\n")  # 上传的文件没有参考摘要
                    f.write(f"INDEX:-1\n")
                
                # 删除临时文件
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                
                return jsonify({
                    "status": "success",
                    "article": text_content,
                    "message": "文件上传成功并已提取内容"
                })
            except Exception as e:
                return jsonify({"status": "error", "message": f"处理Word文档时出错: {str(e)}"})
        else:
            return jsonify({"status": "error", "message": "不支持的文件类型，请上传.docx或.doc文件"})
    except Exception as e:
        return jsonify({"status": "error", "message": f"上传文件时出错: {str(e)}"})

@app.route('/random_news', methods=['GET'])
def random_news():
    try:
        # 使用新闻数据集路径
        dataset_path = "./datasets/news_chinese_simplified_XLSum_v2.0/chinese_simplified_test.jsonl"
        abs_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", dataset_path))
        
        with open(abs_path, 'r', encoding='utf-8') as f:
            test_data = [json.loads(line) for line in f]
        
        # 随机选择一个测试样本
        example_index = random.randint(0, len(test_data) - 1)
        article_text = clean_sentence_spaces(test_data[example_index]["text"])
        reference_summary = clean_sentence_spaces(test_data[example_index]["summary"])
        
        # 保存当前选择的新闻文章和参考摘要
        current_data["news"]["article"] = article_text
        current_data["news"]["reference"] = reference_summary
        current_data["news"]["index"] = example_index
        
        # 将数据写入文件供其他脚本使用
        with open(os.path.join(os.path.dirname(__file__), "current_news.txt"), "w", encoding="utf-8") as f:
            f.write(f"ARTICLE:{article_text}\n")
            f.write(f"REFERENCE:{reference_summary}\n")
            f.write(f"INDEX:{example_index}\n")
        
        return jsonify({
            "status": "success",
            "article": article_text,
            "reference": reference_summary,
            "index": example_index
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/random_media', methods=['GET'])
def random_media():
    try:
        # 使用社交媒体数据集路径
        dataset_path = "./datasets/media_LCSTS/test_public.jsonl"
        abs_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", dataset_path))
        
        with open(abs_path, 'r', encoding='utf-8') as f:
            test_data = [json.loads(line) for line in f]
        
        # 随机选择一个测试样本
        example_index = random.randint(0, len(test_data) - 1)
        article_text = clean_sentence_spaces(test_data[example_index]["text"])
        
        # 保存当前选择的媒体文章
        current_data["media"]["article"] = article_text
        current_data["media"]["reference"] = ""  # 清空参考摘要
        current_data["media"]["index"] = example_index
        
        # 将数据写入文件供其他脚本使用
        with open(os.path.join(os.path.dirname(__file__), "current_media.txt"), "w", encoding="utf-8") as f:
            f.write(f"ARTICLE:{article_text}\n")
            f.write(f"INDEX:{example_index}\n")
        
        return jsonify({
            "status": "success",
            "article": article_text,
            "index": example_index
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/random_encyclopedia', methods=['GET'])
def random_encyclopedia():
    try:
        # 使用百科数据集路径
        dataset_path = "./datasets/baike_wiki/train-00000-of-00001.parquet"
        import pandas as pd
        import numpy as np
        abs_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", dataset_path))
        
        # 读取parquet文件
        df = pd.read_parquet(abs_path)
        
        # 随机选择一个样本
        example_index = random.randint(0, len(df) - 1)
        sample = df.iloc[example_index]
        
        # 从article字典中提取document和summary
        article_dict = sample['article']
        if isinstance(article_dict, dict):
            document = article_dict.get('document', '')
            summary = article_dict.get('summary', '')
            
            # 处理NumPy数组类型
            if isinstance(document, np.ndarray):
                document = ' '.join([str(item) for item in document])
            elif isinstance(document, list):
                document = ' '.join([str(item) for item in document])
                
            if isinstance(summary, np.ndarray):
                summary = ' '.join([str(item) for item in summary])
            elif isinstance(summary, list):
                summary = ' '.join([str(item) for item in summary])
            
            # 确保document和summary是字符串类型
            document = str(document) if document is not None else ''
            summary = str(summary) if summary is not None else ''
                
            article_text = clean_sentence_spaces(document)
            reference_summary = clean_sentence_spaces(summary) if summary else None
            
            # 如果没有摘要，使用第一段作为参考摘要
            if not reference_summary:
                first_para_end = article_text.find("。") + 1
                reference_summary = article_text[:first_para_end] if first_para_end > 0 else article_text[:100]
            
            # 保存当前选择的百科文章和参考摘要
            current_data["encyclopedia"]["article"] = article_text
            current_data["encyclopedia"]["reference"] = reference_summary
            current_data["encyclopedia"]["index"] = example_index
            
            # 将数据写入文件供其他脚本使用
            with open(os.path.join(os.path.dirname(__file__), "current_encyclopedia.txt"), "w", encoding="utf-8") as f:
                f.write(f"ARTICLE:{article_text}\n")
                f.write(f"REFERENCE:{reference_summary}\n")
                f.write(f"INDEX:{example_index}\n")
            
            return jsonify({
                "status": "success",
                "article": article_text,
                "reference": reference_summary,
                "index": example_index
            })
        else:
            return jsonify({"status": "error", "message": f"Article data structure is not as expected: {type(article_dict)}"})
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        return jsonify({"status": "error", "message": f"{str(e)}\n{error_details}"})

# News summarization routes
@app.route('/generate_news_mt5', methods=['POST'])
def generate_news_mt5():
    try:
        # 检查是否使用用户输入的文本
        data = request.json if request.is_json else {}
        use_user_input = data.get('use_user_input', False)
        user_text = data.get('user_text')
        
        if use_user_input and user_text:
            # 当用户输入自己的文本时，参考摘要应该为空
            reference_value = ""
            # 写入 current_news.txt
            current_news_path = os.path.join(os.path.dirname(__file__), "current_news.txt")
            with open(current_news_path, "w", encoding="utf-8") as f:
                f.write(f"ARTICLE:{clean_sentence_spaces(user_text)}\n")
                f.write(f"REFERENCE:{reference_value}\n")
                f.write(f"INDEX:-1\n")
            # 更新全局数据
            current_data["news"]["article"] = clean_sentence_spaces(user_text)
            current_data["news"]["reference"] = reference_value
            current_data["news"]["index"] = -1
        else:
            # 如果不是用户输入的文本，确保保留现有的参考摘要
            # 不需要修改 current_news.txt，因为它已经在 random_news 中被正确设置
            pass
        
        # 运行news_mt5.py (Transformer-based)
        script_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "news_mt5.py"))
        result = subprocess.run(
            [sys.executable, script_path, "--web_mode"],
            capture_output=True, 
            text=True, 
            encoding='utf-8'
        )
        
        # 读取生成的摘要文件
        output_file = os.path.join(os.path.dirname(__file__), "news_mt5_summary.txt")
        with open(output_file, "r", encoding="utf-8") as f:
            content = f.read()
        
        mt5_match = re.search(r'MT5:(.*?)$', content, re.DOTALL)
        mt5_summary = clean_sentence_spaces(mt5_match.group(1)) if mt5_match else ""
        
        return jsonify({
            "status": "success", 
            "mt5_summary": mt5_summary,
            "log": result.stdout
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/generate_news_textrank', methods=['POST'])
def generate_news_textrank():
    try:
        # 检查是否使用用户输入的文本
        data = request.json if request.is_json else {}
        use_user_input = data.get('use_user_input', False)
        user_text = data.get('user_text')
        
        if use_user_input and user_text:
            # 当用户输入自己的文本时，确保清空参考摘要
            with open(os.path.join(os.path.dirname(__file__), "news_mt5_summary.txt"), "w", encoding="utf-8") as f:
                f.write(f"ARTICLE:{clean_sentence_spaces(user_text)}\n")
                f.write(f"REFERENCE:\n")  # 确保参考摘要为空
                f.write(f"MT5:\n")  # 空的MT5摘要，TextRank不需要这个
            
            # 同时更新全局数据，确保一致性
            current_data["news"]["article"] = clean_sentence_spaces(user_text)
            current_data["news"]["reference"] = ""
            current_data["news"]["index"] = -1
        else:
            # 如果不是用户输入的文本，确保保留现有的参考摘要
            # 不需要修改 news_mt5_summary.txt，因为它将会在news_mt5中被正确设置
            pass
        
        # 运行news_textrank.py (TextRank-based)
        script_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "news_textrank.py"))
        result = subprocess.run(
            [sys.executable, script_path, "--web_mode"],
            capture_output=True, 
            text=True, 
            encoding='utf-8'
        )
        
        # 读取生成的摘要文件
        output_file = os.path.join(os.path.dirname(__file__), "news_textrank_results.txt")
        with open(output_file, "r", encoding="utf-8") as f:
            content = f.read()
        
        textrank_match = re.search(r'TextRank:(.*?)(?:\n\n|$)', content, re.DOTALL)
        textrank_summary = clean_sentence_spaces(textrank_match.group(1)) if textrank_match else ""
        
        return jsonify({
            "status": "success", 
            "textrank_summary": textrank_summary,
            "log": result.stdout
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

# News evaluation routes
@app.route('/evaluate_news_mt5', methods=['POST'])
def evaluate_news_mt5():
    try:
        # 评估MT5生成的摘要
        output_file = os.path.join(os.path.dirname(__file__), "news_mt5_summary.txt")
        if not os.path.exists(output_file):
            return jsonify({
                "status": "error",
                "message": "未找到MT5摘要文件，请先生成摘要"
            })
        
        # 使用评估模块进行评估
        results = evaluate_from_file(output_file, 'mt5')
        
        # 返回评估结果和图表URL
        chart_url = url_for('static', filename='summary_eval_chart.png') + f'?t={random.randint(1, 10000)}'
        
        return jsonify({
            "status": "success",
            "evaluation": results,
            "chart_url": chart_url
        })
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"评估MT5摘要时出错: {str(e)}"
        })

@app.route('/evaluate_news_textrank', methods=['POST'])
def evaluate_news_textrank():
    try:
        # 评估TextRank生成的摘要
        output_file = os.path.join(os.path.dirname(__file__), "news_textrank_results.txt")
        if not os.path.exists(output_file):
            return jsonify({
                "status": "error",
                "message": "未找到TextRank摘要文件，请先生成摘要"
            })
        
        # 使用评估模块进行评估
        results = evaluate_from_file(output_file, 'textrank')
        
        # 返回评估结果和图表URL
        chart_url = url_for('static', filename='summary_eval_chart.png') + f'?t={random.randint(1, 10000)}'
        
        return jsonify({
            "status": "success",
            "evaluation": results,
            "chart_url": chart_url
        })
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"评估TextRank摘要时出错: {str(e)}"
        })

# Media summarization routes
@app.route('/generate_media_mt5', methods=['POST'])
def generate_media_mt5():
    try:
        # 检查是否使用用户输入的文本
        data = request.json if request.is_json else {}
        use_user_input = data.get('use_user_input', False)
        user_text = data.get('user_text')
        
        if use_user_input and user_text:
            # 保存用户输入的文本到文件
            with open(os.path.join(os.path.dirname(__file__), "current_media.txt"), "w", encoding="utf-8") as f:
                f.write(f"ARTICLE:{clean_sentence_spaces(user_text)}\n")
                f.write(f"INDEX:-1\n")
            
            # 更新全局数据
            current_data["media"]["article"] = clean_sentence_spaces(user_text)
            current_data["media"]["index"] = -1
        
        # 运行media_mt5.py (Transformer-based)
        script_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "media_mt5.py"))
        result = subprocess.run(
            [sys.executable, script_path, "--web_mode"],
            capture_output=True, 
            text=True, 
            encoding='utf-8'
        )
        
        # 读取生成的摘要文件
        output_file = os.path.join(os.path.dirname(__file__), "media_mt5_summary.txt")
        with open(output_file, "r", encoding="utf-8") as f:
            content = f.read()
        
        mt5_match = re.search(r'MT5:(.*?)$', content, re.DOTALL)
        mt5_summary = clean_sentence_spaces(mt5_match.group(1)) if mt5_match else ""
        
        return jsonify({
            "status": "success", 
            "mt5_summary": mt5_summary,
            "log": result.stdout
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/generate_media_textrank', methods=['POST'])
def generate_media_textrank():
    try:
        # 检查是否使用用户输入的文本
        data = request.json if request.is_json else {}
        use_user_input = data.get('use_user_input', False)
        user_text = data.get('user_text')
        
        if use_user_input and user_text:
            # 保存用户输入的文本到文件 - 确保media_mt5.py已经运行过
            with open(os.path.join(os.path.dirname(__file__), "media_mt5_summary.txt"), "w", encoding="utf-8") as f:
                f.write(f"ARTICLE:{clean_sentence_spaces(user_text)}\n")
                f.write(f"MT5:\n")  # 空的MT5摘要，TextRank不需要这个
        
        # 运行media_textrank.py (TextRank-based)
        script_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "media_textrank.py"))
        result = subprocess.run(
            [sys.executable, script_path, "--web_mode"],
            capture_output=True, 
            text=True, 
            encoding='utf-8'
        )
        
        # 读取生成的摘要文件
        output_file = os.path.join(os.path.dirname(__file__), "media_textrank_results.txt")
        with open(output_file, "r", encoding="utf-8") as f:
            content = f.read()
        
        textrank_match = re.search(r'TextRank:(.*?)(?:\n\n|$)', content, re.DOTALL)
        textrank_summary = clean_sentence_spaces(textrank_match.group(1)) if textrank_match else ""
        
        return jsonify({
            "status": "success", 
            "textrank_summary": textrank_summary,
            "log": result.stdout
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

# Encyclopedia summarization routes
@app.route('/generate_encyclopedia_mt5', methods=['POST'])
def generate_encyclopedia_mt5():
    try:
        # 检查是否使用用户输入的文本
        data = request.json if request.is_json else {}
        use_user_input = data.get('use_user_input', False)
        user_text = data.get('user_text')
        
        if use_user_input and user_text:
            # 保存用户输入的文本到文件
            with open(os.path.join(os.path.dirname(__file__), "current_encyclopedia.txt"), "w", encoding="utf-8") as f:
                f.write(f"ARTICLE:{clean_sentence_spaces(user_text)}\n")
                f.write(f"REFERENCE:\n")  # 用户输入没有参考摘要
                f.write(f"INDEX:-1\n")
            
            # 更新全局数据
            current_data["encyclopedia"]["article"] = clean_sentence_spaces(user_text)
            current_data["encyclopedia"]["reference"] = ""
            current_data["encyclopedia"]["index"] = -1
        
        # 运行encyclopedia_mt5.py (Transformer-based)
        script_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "encyclopedia_mt5.py"))
        result = subprocess.run(
            [sys.executable, script_path, "--web_mode"],
            capture_output=True, 
            text=True, 
            encoding='utf-8'
        )
        
        # 读取生成的摘要文件
        output_file = os.path.join(os.path.dirname(__file__), "encyclopedia_mt5_summary.txt")
        with open(output_file, "r", encoding="utf-8") as f:
            content = f.read()
        
        mt5_match = re.search(r'MT5:(.*?)$', content, re.DOTALL)
        mt5_summary = clean_sentence_spaces(mt5_match.group(1)) if mt5_match else ""
        
        return jsonify({
            "status": "success", 
            "mt5_summary": mt5_summary,
            "log": result.stdout
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/generate_encyclopedia_textrank', methods=['POST'])
def generate_encyclopedia_textrank():
    try:
        # 检查是否使用用户输入的文本
        data = request.json if request.is_json else {}
        use_user_input = data.get('use_user_input', False)
        user_text = data.get('user_text')
        
        if use_user_input and user_text:
            # 保存用户输入的文本到文件 - 确保encyclopedia_mt5.py已经运行过
            with open(os.path.join(os.path.dirname(__file__), "encyclopedia_mt5_summary.txt"), "w", encoding="utf-8") as f:
                f.write(f"ARTICLE:{clean_sentence_spaces(user_text)}\n")
                f.write(f"REFERENCE:\n")  # 用户输入没有参考摘要
                f.write(f"MT5:\n")  # 空的MT5摘要，TextRank不需要这个
        
        # 运行encyclopedia_textrank.py (TextRank-based)
        script_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "encyclopedia_textrank.py"))
        result = subprocess.run(
            [sys.executable, script_path, "--web_mode"],
            capture_output=True, 
            text=True, 
            encoding='utf-8'
        )
        
        # 读取生成的摘要文件
        output_file = os.path.join(os.path.dirname(__file__), "encyclopedia_textrank_results.txt")
        with open(output_file, "r", encoding="utf-8") as f:
            content = f.read()
        
        textrank_match = re.search(r'TextRank摘要:(.*?)(?:\n\n|$)', content, re.DOTALL)
        textrank_summary = clean_sentence_spaces(textrank_match.group(1)) if textrank_match else ""
        
        return jsonify({
            "status": "success", 
            "textrank_summary": textrank_summary,
            "log": result.stdout
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

if __name__ == '__main__':
    # 预先启动一个线程来加载模型，这样第一次批量处理时就不会太慢
    threading.Thread(target=load_mt5_model).start()
    app.run(debug=True, port=5000)